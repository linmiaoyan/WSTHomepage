"""
模板管理API
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
from fastapi.responses import FileResponse, Response
from datetime import datetime
from pathlib import Path
import shutil
from app.database import get_db
from app.models import Template
from app.services.file_handler import extract_placeholders
from app.deps import require_admin
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
import config

router = APIRouter(
    prefix="/api/templates",
    tags=["模板管理"],
    dependencies=[Depends(require_admin)],
)


class TemplateCreate(BaseModel):
    name: str
    description: Optional[str] = None
    created_by: Optional[str] = None


class TemplateResponse(BaseModel):
    id: int
    name: str
    description: Optional[str]
    file_path: str
    file_type: Optional[str]
    placeholders: List[str]
    placeholder_positions: Optional[List[Dict]] = []
    created_by: Optional[str]
    created_at: datetime

    class Config:
        from_attributes = True


@router.get("/", response_model=List[TemplateResponse])
def get_templates(db: Session = Depends(get_db)):
    """获取模板列表"""
    templates = db.query(Template).all()
    # 确保placeholder_positions不为None
    for template in templates:
        if template.placeholder_positions is None:
            template.placeholder_positions = []
    return templates


@router.get("/available-fields")
def get_available_fields():
    """
    获取所有可用的字段名列表
    用于在编辑模板时显示可插入的变量
    """
    # 基础字段（Teacher模型中的字段）
    base_fields = [
        {"name": "name", "label": "姓名", "category": "基本信息"},
        {"name": "sex", "label": "性别", "category": "基本信息"},
        {"name": "id_number", "label": "身份证号", "category": "基本信息"},
        {"name": "phone", "label": "手机号", "category": "联系方式"},
        {"name": "email", "label": "邮箱", "category": "联系方式"},
        {"name": "department", "label": "部门", "category": "工作信息"},
        {"name": "position", "label": "职位", "category": "工作信息"},
        {"name": "title", "label": "职称", "category": "工作信息"},
    ]
    
    # 扩展字段（extra_data中常见的字段）
    extra_fields = [
        {"name": "age", "label": "年龄", "category": "基本信息"},
        {"name": "birth_date", "label": "出生日期", "category": "基本信息"},
        {"name": "education", "label": "学历", "category": "教育背景"},
        {"name": "degree", "label": "学位", "category": "教育背景"},
        {"name": "school", "label": "毕业学校", "category": "教育背景"},
        {"name": "major", "label": "所学专业", "category": "教育背景"},
        {"name": "work_date", "label": "参加工作时间", "category": "工作信息"},
        {"name": "work_years", "label": "工龄", "category": "工作信息"},
        {"name": "teach_years", "label": "教龄", "category": "工作信息"},
        {"name": "political_status", "label": "政治面貌", "category": "其他信息"},
        {"name": "nationality", "label": "民族", "category": "其他信息"},
        {"name": "native_place", "label": "籍贯", "category": "其他信息"},
    ]
    
    return {
        "base_fields": base_fields,
        "extra_fields": extra_fields,
        "all_fields": base_fields + extra_fields
    }


@router.get("/{template_id}", response_model=TemplateResponse)
def get_template(template_id: int, db: Session = Depends(get_db)):
    """获取单个模板"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 直接返回模板对象，Pydantic会自动处理序列化
    # 确保placeholders不为None
    if template.placeholders is None:
        template.placeholders = []
    return template


@router.post("/", response_model=TemplateResponse)
async def create_template(
    name: str,
    description: Optional[str] = None,
    created_by: Optional[str] = None,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """上传模板文件"""
    # 检查文件类型
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in config.ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_ext}")
    
    # 检查文件大小
    # 读取文件内容以检查大小
    file_content = await file.read()
    file_size = len(file_content)
    
    # HTML文件使用10MB限制，其他文件使用50MB限制
    if file_ext == '.html':
        max_size = config.MAX_HTML_FILE_SIZE
        max_size_mb = 10
    else:
        max_size = config.MAX_FILE_SIZE
        max_size_mb = 50
    
    if file_size > max_size:
        raise HTTPException(
            status_code=413, 
            detail=f"文件大小超过限制：{file_size / (1024 * 1024):.2f}MB，最大允许{max_size_mb}MB"
        )
    
    # 保存文件
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = "".join(c for c in file.filename if c.isalnum() or c in "._-")
    file_path = config.TEMPLATE_DIR / f"{timestamp}_{safe_filename}"
    
    with open(file_path, "wb") as buffer:
        buffer.write(file_content)
    
    # PDF文件不需要自动提取占位符，由用户在界面上拖动选择位置
    placeholders = []
    if file_ext == '.pdf':
        # PDF文件，占位符位置由用户手动设置
        pass
    else:
        # 其他文件类型（已废弃，但保留兼容性）
        placeholders = extract_placeholders(str(file_path))
    
    # 创建模板记录
    template = Template(
        name=name,
        description=description,
        file_path=str(file_path),
        file_type=file_ext,
        placeholders=placeholders,
        placeholder_positions=[],  # 初始为空，由用户设置
        created_by=created_by
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return template


@router.put("/{template_id}/placeholders", response_model=TemplateResponse)
def update_placeholders(
    template_id: int,
    placeholders: List[str],
    db: Session = Depends(get_db)
):
    """更新占位符列表"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    template.placeholders = placeholders
    db.commit()
    db.refresh(template)
    return template


@router.get("/{template_id}/content")
def get_template_content(template_id: int, db: Session = Depends(get_db)):
    """获取模板文件内容，用于在线编辑"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = Path(template.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="模板文件不存在")
    
    file_ext = file_path.suffix.lower()
    
    if file_ext in ['.docx', '.doc']:
        # 读取Word文档内容
        from docx import Document
        doc = Document(file_path)
        
        # 提取段落内容
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 提取表格内容
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append({
                    "index": table_idx,
                    "rows": table_data
                })
        
        return {
            "template_id": template_id,
            "file_type": "docx",
            "paragraphs": paragraphs,
            "tables": tables
        }
    
    elif file_ext in ['.xlsx', '.xls']:
        # 读取Excel文件内容
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=False)
        
        sheets = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = []
            
            # 获取有数据的范围
            if sheet.max_row > 0 and sheet.max_column > 0:
                for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, 
                                         min_col=1, max_col=sheet.max_column, values_only=False):
                    row_data = []
                    for cell in row:
                        # 获取单元格值，如果是公式则保留公式
                        if cell.value is None:
                            row_data.append("")
                        else:
                            row_data.append(str(cell.value))
                    sheet_data.append(row_data)
            
            sheets.append({
                "name": sheet_name,
                "data": sheet_data
            })
        
        return {
            "template_id": template_id,
            "file_type": "xlsx",
            "sheets": sheets
        }
    
    elif file_ext == '.pdf':
        # PDF文件：返回PDF信息和页面数量
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            
            # 获取第一页尺寸
            if num_pages > 0:
                page = reader.pages[0]
                page_width = float(page.mediabox.width)
                page_height = float(page.mediabox.height)
            else:
                page_width = 595
                page_height = 842
            
            return {
                "template_id": template_id,
                "file_type": "pdf",
                "num_pages": num_pages,
                "page_width": page_width,
                "page_height": page_height
            }
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF处理库未安装，请安装: pip install PyPDF2")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"读取PDF失败: {str(e)}")
    
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_ext}")


class TemplateContentUpdate(BaseModel):
    file_type: str
    paragraphs: Optional[List[str]] = None
    tables: Optional[List[Dict]] = None
    sheets: Optional[List[Dict]] = None


@router.put("/{template_id}/content")
def update_template_content(
    template_id: int,
    content: TemplateContentUpdate,
    db: Session = Depends(get_db)
):
    """保存编辑后的模板文件内容"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = Path(template.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="模板文件不存在")
    
    file_ext = file_path.suffix.lower()
    file_type = content.file_type
    
    try:
        if file_type == "docx" and file_ext in ['.docx', '.doc']:
            # 保存Word文档
            from docx import Document
            doc = Document()
            
            # 添加段落
            paragraphs = content.paragraphs or []
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
            
            # 添加表格
            tables = content.tables or []
            for table_data in tables:
                rows = table_data.get("rows", [])
                if rows:
                    # 确定列数
                    max_cols = max(len(row) for row in rows) if rows else 1
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < max_cols:
                                table.rows[row_idx].cells[col_idx].text = str(cell_text) if cell_text else ""
            
            doc.save(file_path)
        
        elif file_type == "xlsx" and file_ext in ['.xlsx', '.xls']:
            # 保存Excel文件
            from openpyxl import Workbook
            wb = Workbook()
            wb.remove(wb.active)  # 删除默认工作表
            
            sheets = content.sheets or []
            for sheet_data in sheets:
                sheet_name = sheet_data.get("name", "Sheet1")
                sheet = wb.create_sheet(title=sheet_name)
                
                rows = sheet_data.get("data", [])
                for row_idx, row_data in enumerate(rows, start=1):
                    for col_idx, cell_value in enumerate(row_data, start=1):
                        if cell_value:
                            sheet.cell(row=row_idx, column=col_idx, value=cell_value)
            
            wb.save(file_path)
        
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_type}")
        
        # 重新提取占位符
        from app.services.file_handler import extract_placeholders
        placeholders = extract_placeholders(str(file_path))
        template.placeholders = placeholders
        db.commit()
        
        return {"message": "保存成功", "placeholders": placeholders}
    
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")


@router.get("/{template_id}/pdf-file")
def get_pdf_file(template_id: int, db: Session = Depends(get_db)):
    """获取PDF文件（用于浏览器内嵌显示）"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = Path(template.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="模板文件不存在")
    
    if template.file_type != '.pdf':
        raise HTTPException(status_code=400, detail="此模板不是PDF文件")
    
    from fastapi.responses import FileResponse
    # 不设置filename，让浏览器内嵌显示而不是下载
    return FileResponse(
        path=str(file_path),
        media_type="application/pdf",
        headers={"Content-Disposition": "inline"}  # inline表示内嵌显示，而不是attachment下载
    )


@router.get("/{template_id}/pdf-preview/{page}")
def get_pdf_preview(template_id: int, page: int, db: Session = Depends(get_db)):
    """获取PDF页面的预览图片（备用方案，如果PDF.js不可用）"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = Path(template.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="模板文件不存在")
    
    if template.file_type != '.pdf':
        raise HTTPException(status_code=400, detail="此模板不是PDF文件")
    
    try:
        from app.services.pdf_handler import get_pdf_preview
        preview_image = get_pdf_preview(str(file_path), page)
        return Response(content=preview_image, media_type="image/png")
    except Exception as e:
        # 即使预览失败，也返回一个占位图片，避免前端报错
        from io import BytesIO
        try:
            from PIL import Image, ImageDraw
            img = Image.new('RGB', (595, 842), color='white')
            draw = ImageDraw.Draw(img)
            text = f"PDF预览需要安装pdf2image和poppler\n\n请运行：\npip install pdf2image\n\n并安装poppler工具"
            draw.text((50, 350), text, fill='red')
            img_byte_arr = BytesIO()
            img.save(img_byte_arr, format='PNG')
            return Response(content=img_byte_arr.getvalue(), media_type="image/png")
        except ImportError:
            # 如果PIL也不可用，返回错误
            raise HTTPException(status_code=500, detail="PDF预览功能不可用，请使用PDF.js直接显示")


@router.put("/{template_id}/placeholder-positions")
def update_placeholder_positions(
    template_id: int,
    positions: List[Dict[str, Any]],
    db: Session = Depends(get_db)
):
    """保存占位符位置信息"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 验证位置数据格式
    for pos in positions:
        if "field_name" not in pos or "page" not in pos or "x" not in pos or "y" not in pos:
            raise HTTPException(status_code=400, detail="位置信息格式错误，必须包含field_name, page, x, y")
    
    # 确保is_signature字段被正确保存（转换为布尔值）
    for pos in positions:
        if "is_signature" in pos:
            # 确保is_signature是布尔值
            is_sig = pos.get("is_signature")
            if isinstance(is_sig, str):
                pos["is_signature"] = is_sig.lower() in ['true', '1', 'yes']
            elif isinstance(is_sig, int):
                pos["is_signature"] = bool(is_sig)
            # 如果已经是布尔值，保持不变
    
    # 更新占位符位置
    template.placeholder_positions = positions
    
    # 提取占位符列表（从位置信息中提取）
    placeholders = list(set([pos.get("field_name") for pos in positions if pos.get("field_name")]))
    template.placeholders = placeholders
    
    db.commit()
    db.refresh(template)
    
    return {"message": "保存成功", "placeholders": placeholders}


@router.get("/{template_id}/related-tasks")
def get_template_related_tasks(template_id: int, db: Session = Depends(get_db)):
    """获取模板关联的任务列表"""
    from app.models import Task
    
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    related_tasks = db.query(Task).filter(Task.template_id == template_id).all()
    return {
        "template_id": template_id,
        "template_name": template.name,
        "related_tasks": [
            {
                "id": task.id,
                "name": task.name,
                "status": task.status,
                "created_at": task.created_at.isoformat() if task.created_at else None
            }
            for task in related_tasks
        ],
        "count": len(related_tasks)
    }


@router.delete("/{template_id}")
def delete_template(template_id: int, db: Session = Depends(get_db)):
    """删除模板，同时删除关联的任务"""
    import os
    from app.models import Task
    
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 查询关联的任务
    related_tasks = db.query(Task).filter(Task.template_id == template_id).all()
    deleted_tasks = []
    
    # 删除所有关联的任务及其导出文件
    for task in related_tasks:
        # 删除任务的导出文件（如果存在）
        if task.export_path:
            export_path = Path(task.export_path)
            if export_path.exists():
                try:
                    os.remove(export_path)
                except Exception as e:
                    print(f"删除任务导出文件失败: {e}")
        
        deleted_tasks.append(task.name)
        db.delete(task)
    
    # 删除模板文件
    if Path(template.file_path).exists():
        try:
            Path(template.file_path).unlink()
        except Exception as e:
            print(f"删除模板文件失败: {e}")
    
    # 删除模板记录
    db.delete(template)
    db.commit()
    
    message = "删除成功"
    if deleted_tasks:
        message += f"，已同时删除 {len(deleted_tasks)} 个关联任务：{', '.join(deleted_tasks)}"
    
    return {"message": message, "deleted_tasks": deleted_tasks}
