"""
模板处理服务
"""
import os
import re
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
import shutil
import config


def fill_docx_template(template_path: str, data: Dict[str, Any], output_path: str):
    """
    填充Word模板
    data: 教师数据字典
    """
    # 复制模板文件
    shutil.copy(template_path, output_path)
    doc = Document(output_path)
    
    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        paragraph.text = replace_placeholders(paragraph.text, data)
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_placeholders(cell.text, data)
    
    doc.save(output_path)


def fill_xlsx_template(template_path: str, data: Dict[str, Any], output_path: str):
    """
    填充Excel模板
    """
    # 复制模板文件
    shutil.copy(template_path, output_path)
    wb = load_workbook(output_path)
    
    # 遍历所有工作表
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        
        # 遍历所有单元格
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = replace_placeholders(cell.value, data)
    
    wb.save(output_path)


def replace_placeholders(text: str, data: Dict[str, Any]) -> str:
    """
    替换文本中的占位符
    支持 {{field_name}} 格式
    也支持从 extra_data 中获取数据
    """
    def replace_func(match):
        field_name = match.group(1)
        # 先从主字段查找
        if field_name in data:
            value = data[field_name]
        # 再从extra_data查找
        elif 'extra_data' in data and isinstance(data['extra_data'], dict):
            value = data['extra_data'].get(field_name, '')
        else:
            value = ''
        
        # 处理None值
        if value is None:
            value = ''
        
        return str(value)
    
    return re.sub(r'\{\{(\w+)\}\}', replace_func, text)


def process_template(template_path: str, data: Dict[str, Any], output_path: str, placeholder_positions: List[Dict[str, Any]] = None):
    """
    根据文件类型处理模板
    
    Args:
        template_path: 模板文件路径
        data: 教师数据字典
        output_path: 输出文件路径
        placeholder_positions: 占位符位置信息（仅PDF需要）
    """
    print(f"[模板处理] 开始处理模板: {template_path}")
    print(f"[模板处理] 输出路径: {output_path}")
    print(f"[模板处理] 数据字段数量: {len(data)}")
    
    # 检查是否有签名字段
    signature_fields = []
    if 'extra_data' in data and isinstance(data['extra_data'], dict):
        for key, value in data['extra_data'].items():
            if isinstance(value, str) and value.startswith('data:image'):
                signature_fields.append(key)
                print(f"[模板处理] 检测到签名字段: {key}, 数据长度: {len(value)} 字符")
    
    ext = Path(template_path).suffix.lower()
    print(f"[模板处理] 文件类型: {ext}")
    
    try:
        if ext == '.pdf':
            # PDF文件：在指定位置添加文本
            print(f"[模板处理] 处理PDF模板...")
            from app.services.pdf_handler import add_text_to_pdf
            if not placeholder_positions:
                raise ValueError("PDF模板需要提供占位符位置信息")
            print(f"[模板处理] 占位符数量: {len(placeholder_positions)}")
            print(f"[模板处理] 调用 add_text_to_pdf...")
            add_text_to_pdf(template_path, output_path, placeholder_positions, data)
            print(f"[模板处理] PDF处理完成")
        elif ext in ['.docx', '.doc']:
            print(f"[模板处理] 处理Word模板...")
            fill_docx_template(template_path, data, output_path)
            print(f"[模板处理] Word处理完成")
        elif ext in ['.xlsx', '.xls']:
            print(f"[模板处理] 处理Excel模板...")
            fill_xlsx_template(template_path, data, output_path)
            print(f"[模板处理] Excel处理完成")
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
        print(f"[模板处理] 模板处理成功: {output_path}")
    except Exception as e:
        print(f"[模板处理] 模板处理失败: {e}")
        import traceback
        traceback.print_exc()
        raise

